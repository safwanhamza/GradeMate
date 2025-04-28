# drive_processor.py  ── safe UTF-8 + logging setup
import os
import sys
import io
import time
import logging
from datetime import datetime
from typing import List, Dict, Any, Optional

# ──────────────────────────────────────────────────────────────────────────────
# 1. Timestamp helper – use Django's timezone.now() if available, else datetime
try:
    from django.utils.timezone import now            # type: ignore
except Exception:                                     # pragma: no cover
    now = datetime.now                                # noqa: N816  (function, not constant)

# ──────────────────────────────────────────────────────────────────────────────
# 2. Force UTF-8 console I/O on Windows **without** detaching the buffer
if sys.platform == "win32" and hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8",
                           errors="replace",
                           line_buffering=True)
    sys.stderr.reconfigure(encoding="utf-8",
                           errors="replace",
                           line_buffering=True)

# ──────────────────────────────────────────────────────────────────────────────
# 3. Logging – one file + stdout
LOG_FILE = f"drive_pipeline_{now().strftime('%Y%m%d_%H%M%S')}.log"

logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILE, encoding="utf-8"),
        logging.StreamHandler(stream=sys.stdout),
    ],
)
logger = logging.getLogger(__name__)

# ──────────────────────────────────────────────────────────────────────────────
# 4. Third-party / project imports (unchanged)
from unstructured.partition.auto import partition
from unstructured.documents.elements import (
    Image as UnstructuredImage,
    Table as UnstructuredTable,
    Text,
    Title,
    ListItem,
    NarrativeText,
    Element,
)
from langchain.schema import Document
from .models import ChunkedText, DrivePDF

# Directory paths
DOWNLOAD_DIR = "downloaded_files"
FIGURES_DIR = "extracted_figures"
TEXT_TYPES = (Text, Title, ListItem, NarrativeText)

# Flag controlling Django DB use
USE_DJANGO_DB_STORAGE = True



def verify_embedding_model(embedding_model) -> bool:
    """
    Simple sanity-check for any embedding model.
    """
    test_text = "This is a test query to verify the embedding model is working"
    try:
        if hasattr(embedding_model, "embed_query"):
            vec = embedding_model.embed_query(test_text)
        elif hasattr(embedding_model, "embed_documents"):
            vec = embedding_model.embed_documents([test_text])[0]
        else:
            vec = embedding_model.embed(test_text)          # type: ignore
        return bool(vec and isinstance(vec, list))
    except Exception:
        return False

def handle_fallback_file(path):
    """
    Extract text from file when partition fails
    """
    ext = os.path.splitext(path)[1].lower()
    logger.info(f"Using fallback handler for file type: {ext}")
    
    try:
        # Handle PDF files
        if ext == '.pdf':
            from PyPDF2 import PdfReader
            reader = PdfReader(path)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            logger.info(f"Extracted {len(text)} characters from PDF")
            return [NarrativeText(text=text)]
        
        # Handle CSV files
        elif ext == '.csv':
            import pandas as pd
            df = pd.read_csv(path)
            metadata = f"CSV with {len(df)} rows and {len(df.columns)} columns."
            sample = df.head(3).to_string()
            logger.info(f"Processed CSV: {len(df)} rows, {len(df.columns)} columns")
            return [NarrativeText(text=f"{metadata}\n\nSample data:\n{sample}")]
        
        # Handle HTML files
        elif ext in ['.html', '.htm']:
            from bs4 import BeautifulSoup
            with open(path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f, 'html.parser')
            title = soup.title.string if soup.title else "No title"
            body = soup.get_text(separator='\n', strip=True)
            logger.info(f"Processed HTML: {len(body)} characters")
            return [
                Title(text=f"HTML Title: {title}"),
                NarrativeText(text=body)
            ]
        
        # Handle Python files
        elif ext == '.py':
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            logger.info(f"Processed Python file: {len(content)} characters")
            return [NarrativeText(text=content)]
        
        # Handle Jupyter notebooks
        elif ext == '.ipynb':
            import json
            with open(path, 'r', encoding='utf-8') as f:
                nb = json.load(f)
            elements = []
            for cell in nb.get('cells', []):
                if 'source' in cell:
                    text = ''.join(cell.get('source', []))
                    cell_type = cell.get('cell_type', 'unknown')
                    elements.append(NarrativeText(text=f"[{cell_type.upper()}]\n{text}"))
            logger.info(f"Processed Jupyter notebook: {len(elements)} cells")
            return elements if elements else [NarrativeText(text="Empty notebook or no valid content found.")]
        
        # Handle other text-based files
        elif ext in ['.txt', '.md', '.json', '.xml']:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            logger.info(f"Processed text file: {len(content)} characters")
            return [NarrativeText(text=content)]
        
        # Handle Word documents
        elif ext in ['.docx', '.doc']:
            try:
                from docx import Document
                doc = Document(path)
                text = "\n".join([para.text for para in doc.paragraphs])
                logger.info(f"Processed Word document: {len(text)} characters")
                return [NarrativeText(text=text)]
            except ImportError:
                logger.warning("python-docx not installed. Cannot process Word document.")
                return [NarrativeText(text=f"Unable to process Word document: {os.path.basename(path)}. python-docx not installed.")]
        
        # Default case
        logger.info(f"Used generic handler for file: {path}")
        return [NarrativeText(text=f"File content from: {os.path.basename(path)}")]
    
    except Exception as e:
        logger.exception(f"Fallback processing failed for {path}: {e}")
        return [NarrativeText(text=f"Error processing file: {str(e)}")]

def process_document_simple(path: str, output_figures_path: Optional[str] = None) -> List[Element]:
    """
    Process a document without relying on complex models or external services.
    
    Args:
        path: Path to the document file
        output_figures_path: Directory to store extracted figures (optional)
    
    Returns:
        List of document elements
    """
    file_ext = os.path.splitext(path)[1].lower()
    logger.info(f"Processing {path} with extension {file_ext}")
    
    try:
        # Try to use unstructured's partition function first
        try:
            extract_params = {
                "filename": path,
                "strategy": "hi_res",
                "languages": ["eng"],
            }
            
            # Only enable image extraction if output directory is provided
            if output_figures_path:
                extract_params.update({
                    "extract_image_block_to_payload": True,
                    "extract_image_block_types": ["Image", "Table"],
                    "infer_table_structure": True,
                    "image_output_dir_path": output_figures_path
                })
            
            logger.info(f"Attempting partition with params: {extract_params}")
            elements = partition(**extract_params)
            
            if elements and len(elements) > 0:
                logger.info(f"Successfully extracted {len(elements)} elements with partition")
                # Log element types for debugging
                element_types = {}
                for el in elements:
                    el_type = type(el).__name__
                    element_types[el_type] = element_types.get(el_type, 0) + 1
                logger.info(f"Element breakdown: {element_types}")
                return elements
            else:
                logger.warning("Partition returned no elements, falling back to manual extraction")
                
        except Exception as e:
            logger.warning(f"Partition failed for {path}: {e}")
    
    except ImportError:
        logger.warning("Unstructured library not properly installed")
    
    # Fall back to simple extraction
    logger.info(f"Using fallback processing for {path}")
    return handle_fallback_file(path)

def create_chroma_db(chunks, embedding_model, persist_directory=None, progress_callback=None):
    """
    Create a Chroma vector store from chunks.
    
    Args:
        chunks: List of Document objects to store
        embedding_model: The embedding model to use
        persist_directory: Directory to persist the Chroma DB (optional)
        progress_callback: Callback function for progress updates
    
    Returns:
        tuple: (success, vector_store_or_error_message)
    """
    try:
        from langchain_community.vectorstores import Chroma
        import shutil
        
        # Define the Chroma DB path if not provided
        if not persist_directory:
            persist_directory = os.path.join(os.getcwd(), "chroma_db")
        
        logger.info(f"Creating Chroma DB at {persist_directory} with {len(chunks)} chunks")
        
        # Verify embedding model is working
        embedding_test_success = verify_embedding_model(embedding_model)
        if not embedding_test_success:
            error_msg = "Embedding model test failed, aborting Chroma DB creation"
            logger.error(error_msg)
            if progress_callback:
                progress_callback(f"❌ {error_msg}")
            return False, error_msg
        
        # Remove existing DB if it exists
        if os.path.exists(persist_directory):
            logger.info(f"Removing existing Chroma DB at {persist_directory}")
            if progress_callback:
                progress_callback(f"🗑️ Removing existing Chroma DB at {persist_directory}")
            shutil.rmtree(persist_directory)
        
        # Create directory if it doesn't exist
        os.makedirs(persist_directory, exist_ok=True)
        
        # Create new DB
        logger.info(f"Creating new Chroma DB with {len(chunks)} chunks")
        if progress_callback:
            progress_callback(f"📥 Creating new Chroma DB with {len(chunks)} chunks")
        
        start_time = time.time()
        vector_store = Chroma.from_documents(
            documents=chunks,
            embedding=embedding_model,
            persist_directory=persist_directory
        )
        
        # Make sure to persist the DB
        vector_store.persist()
        end_time = time.time()
        
        logger.info(f"✅ Successfully created Chroma DB in {end_time - start_time:.2f} seconds")
        
        # Verify the DB was created properly with a test query
        try:
            logger.info(f"Verifying DB contents with test query...")
            results = vector_store.similarity_search("test query", k=1)
            logger.info(f"Verification query returned {len(results)} results")
            
            if progress_callback:
                progress_callback(f"✅ Successfully created and verified Chroma DB with {len(chunks)} chunks")
            
            return True, vector_store
        except Exception as e:
            error_msg = f"DB verification failed: {str(e)}"
            logger.error(error_msg)
            if progress_callback:
                progress_callback(f"⚠️ {error_msg}")
            return False, error_msg
        
    except Exception as e:
        error_msg = f"Error creating Chroma DB: {str(e)}"
        logger.exception(error_msg)
        if progress_callback:
            progress_callback(f"❌ {error_msg}")
        import traceback
        logger.error(traceback.format_exc())
        return False, str(e)

def process_drive_pdf(drive_pdf: DrivePDF, embedding_model=None, force_reprocess=False, direct_to_chroma=False) -> Dict[str, Any]:
    """
    Process a document and store chunks in the database and optionally directly to Chroma.
    
    Args:
        drive_pdf: DrivePDF object to process
        embedding_model: Embedding model to use (required if direct_to_chroma=True)
        force_reprocess: Force reprocessing even if chunks already exist
        direct_to_chroma: If True, bypass Django DB and store directly to Chroma
    
    Returns:
        Dictionary with processing results
    """
    logger.info(f"Processing {drive_pdf.title} (direct_to_chroma={direct_to_chroma})")
    
    # Check for direct_to_chroma requirements
    if direct_to_chroma and embedding_model is None:
        error_msg = "Embedding model is required when direct_to_chroma=True"
        logger.error(error_msg)
        return {'status': 'error', 'error': error_msg}
    
    # If using Django DB, check if already processed
    if not direct_to_chroma:
        existing_chunks = ChunkedText.objects.filter(pdf=drive_pdf).count()
        if existing_chunks > 0 and not force_reprocess:
            logger.info(f"Already processed {drive_pdf.title} ({existing_chunks} chunks)")
            return {'status': 'already_processed', 'chunk_count': existing_chunks}
        elif existing_chunks > 0 and force_reprocess:
            logger.info(f"Force reprocessing {drive_pdf.title} (deleting {existing_chunks} existing chunks)")
            # Delete existing chunks to reprocess
            ChunkedText.objects.filter(pdf=drive_pdf).delete()
    
    # Ensure directories exist
    os.makedirs(DOWNLOAD_DIR, exist_ok=True)
    os.makedirs(FIGURES_DIR, exist_ok=True)
    
    # Get local file path
    filename = drive_pdf.title
    local_path = os.path.join(DOWNLOAD_DIR, filename)
    
    # Check if file exists locally
    if not os.path.exists(local_path):
        logger.error(f"File not found: {local_path}")
        return {'status': 'error', 'error': f"File not found: {local_path}"}
    
    try:
        # Create output directory for figures
        output_figures_path = os.path.join(FIGURES_DIR, os.path.splitext(filename)[0])
        os.makedirs(output_figures_path, exist_ok=True)
        
        # Process the document
        elements = process_document_simple(local_path, output_figures_path)
        
        if not elements or len(elements) == 0:
            logger.warning(f"No elements extracted from {filename}")
            return {'status': 'error', 'error': f"No elements extracted from {filename}"}
        
        # Extract text from elements
        texts = []
        for el in elements:
            if isinstance(el, TEXT_TYPES):
                texts.append(el.text)
        
        combined_text = "\n\n".join(texts)
        
        # Simple chunking by paragraphs or fixed length
        chunks = []
        paragraphs = combined_text.split("\n\n")
        
        current_chunk = ""
        for para in paragraphs:
            if len(current_chunk) + len(para) > 1000:  # Limit chunk size
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = para
            else:
                current_chunk += "\n\n" + para if current_chunk else para
        
        if current_chunk:  # Don't forget the last chunk
            chunks.append(current_chunk)
        
        # If no chunks were created, create at least one
        if not chunks and combined_text:
            chunks = [combined_text]
        elif not chunks:
            chunks = [f"Empty or unprocessable file: {filename}"]
        
        logger.info(f"Created {len(chunks)} text chunks from {filename}")
        
        # Option 1: Store in Django DB
        if not direct_to_chroma:
            # Clear previous chunks if any
            ChunkedText.objects.filter(pdf=drive_pdf).delete()
            
            # Save chunks to database
            for i, chunk in enumerate(chunks):
                ChunkedText.objects.create(
                    pdf=drive_pdf,
                    content=chunk,
                    order=i
                )
            
            logger.info(f"Saved {len(chunks)} chunks to Django DB for {filename}")
            
            return {
                'status': 'success', 
                'chunk_count': len(chunks),
                'file': filename
            }
            
        # Option 2: Store directly in Chroma
        else:
            logger.info(f"Preparing {len(chunks)} chunks for direct Chroma storage")
            
            # Convert text chunks to Document objects
            doc_chunks = []
            for i, chunk in enumerate(chunks):
                doc_chunks.append(
                    Document(
                        page_content=chunk,
                        metadata={
                            "source": filename,
                            "pdf_id": drive_pdf.id,
                            "order": i,
                            "type": "text"
                        }
                    )
                )
            
            # Create Chroma DB for this file
            chroma_path = os.path.join(os.getcwd(), "chroma_db", f"file_{drive_pdf.id}")
            success, result = create_chroma_db(
                doc_chunks, 
                embedding_model, 
                persist_directory=chroma_path
            )
            
            if success:
                logger.info(f"Successfully stored {len(chunks)} chunks directly to Chroma for {filename}")
                return {
                    'status': 'success',
                    'chunk_count': len(chunks),
                    'file': filename,
                    'chroma_path': chroma_path
                }
            else:
                logger.error(f"Failed to store chunks in Chroma: {result}")
                return {
                    'status': 'error',
                    'error': f"Chroma storage failed: {result}",
                    'file': filename
                }
            
    except Exception as e:
        logger.exception(f"Processing failed for {filename}: {e}")
        return {
            'status': 'error',
            'error': str(e),
            'file': filename
        }




def process_all_drive_pdfs(user, progress_callback=None, force_reprocess=True):
    """
    Process all DrivePDF objects for a user using Nomic Embeddings.
    
    Args:
        user: User whose files to process
        progress_callback: Optional callback function to report progress
        force_reprocess: Force reprocessing even if chunks might exist
    
    Returns:
        dict: Processing statistics
    """
    import os
    import logging
    import traceback
    from typing import Optional, Any, List
    from langchain_nomic.embeddings import NomicEmbeddings
    from langchain.schema import Document
    from langchain_community.vectorstores import Chroma

    # Set up logging
    logger = logging.getLogger(__name__)

    # Retrieve drive PDFs
    drive_pdfs = DrivePDF.objects.filter(user=user)
    
    if not drive_pdfs.exists():
        if progress_callback:
            progress_callback("No Drive PDFs found to process.")
        return {
            "successful_files": 0,
            "failed_files": [],
            "status": "no_files"
        }
    
    def setup_embedding_model() -> Optional[Any]:
        """
        Set up Nomic Embeddings with robust error handling.
        
        Returns:
            Embedding model or None if setup fails
        """
        try:
            # Ensure Nomic API key is set
            nomic_api_key = os.environ.get('NOMIC_API_KEY')
            if not nomic_api_key:
                # Set the API key directly if not in environment
                nomic_api_key = "nk-rjqDQKJtuoRaTcocIvaSJr6g5JItcyLvNJR4O7h153o"
                os.environ['NOMIC_API_KEY'] = nomic_api_key
            
            # Create Nomic Embeddings
            embedding_model = NomicEmbeddings(
                model="nomic-embed-text-v1"
                # Optionally specify dimensionality or other parameters
            )
            
            # Verify the embedding
            test_text = "This is a test query to verify the Nomic embedding model is working"
            test_embedding = embedding_model.embed_query(test_text)
            
            if not test_embedding or len(test_embedding) == 0:
                logger.error("Nomic Embedding verification failed: Empty embedding vector")
                return None
            
            logger.info("Successfully set up Nomic Embeddings")
            return embedding_model
        
        except Exception as e:
            logger.error(f"Error setting up Nomic Embeddings: {str(e)}")
            logger.error(traceback.format_exc())
            return None

    def process_document(drive_pdf, embedding_model) -> Optional[List[Document]]:
        """
        Process a single document and convert to LangChain Documents
        
        Args:
            drive_pdf: DrivePDF object to process
            embedding_model: Embedding model to use
        
        Returns:
            List of LangChain Documents or None if processing fails
        """
        # Ensure directories exist
        os.makedirs("downloaded_files", exist_ok=True)
        os.makedirs("extracted_figures", exist_ok=True)

        # Construct file paths
        local_path = os.path.join("downloaded_files", drive_pdf.title)
        output_figures_path = os.path.join("extracted_figures", os.path.splitext(drive_pdf.title)[0])
        os.makedirs(output_figures_path, exist_ok=True)

        # Debug logging: Move this before any file operations
        logger.info(f"Processing document: {drive_pdf.title}")
        logger.info(f"Attempting to process file at: {local_path}")

        try:
            # Check if file exists
            if not os.path.exists(local_path):
                logger.error(f"File does not exist: {local_path}")
                return None

            # Try to read the file content
            try:
                with open(local_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            except UnicodeDecodeError:
                # Try reading with a different encoding
                try:
                    with open(local_path, 'r', encoding='latin-1') as f:
                        content = f.read()
                except Exception as e:
                    logger.error(f"Failed to read file {local_path} with any encoding: {e}")
                    return None

            # Simple chunking
            max_chunk_size = 1000
            chunks = []
            
            # Split the content into paragraphs
            paragraphs = content.split('\n\n')
            
            current_chunk = ""
            for para in paragraphs:
                # If adding this paragraph would exceed max chunk size, start a new chunk
                if len(current_chunk) + len(para) > max_chunk_size:
                    if current_chunk:
                        chunks.append(current_chunk)
                    current_chunk = para
                else:
                    # Otherwise, add to current chunk
                    current_chunk += "\n\n" + para if current_chunk else para
            
            # Add the last chunk if not empty
            if current_chunk:
                chunks.append(current_chunk)

            # Log chunk information
            logger.info(f"Created {len(chunks)} chunks for {drive_pdf.title}")

            # Convert chunks to LangChain Documents
            documents = [
                Document(
                    page_content=chunk,
                    metadata={
                        "source": drive_pdf.title,
                        "pdf_id": drive_pdf.id
                    }
                ) for chunk in chunks
            ]

            return documents

        except Exception as e:
            logger.error(f"Unexpected error processing document {drive_pdf.title}: {e}")
            logger.error(traceback.format_exc())
            return None
        










    
    # Setup the embedding model
    embedding_model = setup_embedding_model()
    
    # If no embedding model could be set up, raise an error
    if embedding_model is None:
        error_msg = "Failed to set up Nomic embedding model. Cannot proceed with document processing."
        logger.error(error_msg)
        if progress_callback:
            progress_callback(f"❌ {error_msg}")
        return {
            "successful_files": 0,
            "failed_files": [],
            "status": "embedding_error",
            "error": error_msg
        }
    
    successful_files = 0
    failed_files = []
    all_documents = []
    
    # Process each file
    for drive_pdf in drive_pdfs:
        if progress_callback:
            progress_callback(f"🔄 Processing: {drive_pdf.title}")
        
        try:
            # Process the document
            documents = process_document(drive_pdf, embedding_model)
            
            if documents:
                successful_files += 1
                all_documents.extend(documents)
                
                if progress_callback:
                    progress_callback(f"✅ Successfully processed {drive_pdf.title} with {len(documents)} chunks.")
            else:
                failed_files.append({
                    'title': drive_pdf.title,
                    'error': 'Failed to process document'
                })
                if progress_callback:
                    progress_callback(f"❌ Failed to process {drive_pdf.title}")
        
        except Exception as e:
            logger.exception(f"Unexpected error processing {drive_pdf.title}")
            failed_files.append({
                'title': drive_pdf.title,
                'error': str(e)
            })
            if progress_callback:
                progress_callback(f"❌ Unexpected error processing {drive_pdf.title}: {str(e)}")
    
    # If we have documents, create Chroma vector store
    if all_documents:
        try:
            # Define Chroma DB path
            persist_directory = os.path.join(os.getcwd(), "chroma_db")
            
            # Remove existing Chroma DB if it exists
            if os.path.exists(persist_directory):
                import shutil
                shutil.rmtree(persist_directory)
            
            # Create new Chroma vector store
            Chroma.from_documents(
                documents=all_documents,
                embedding=embedding_model,
                persist_directory=persist_directory
            )

            # NEW: Verify Chroma DB details
            def verify_chroma_db(persist_directory):
                try:
                    from langchain_community.vectorstores import Chroma
                    from langchain_nomic.embeddings import NomicEmbeddings
                    
                    embedding_model = NomicEmbeddings(model="nomic-embed-text-v1")
                    vector_store = Chroma(
                        persist_directory=persist_directory, 
                        embedding_function=embedding_model
                    )
                    
                    total_chunks = vector_store._collection.count()
                    logger.info(f"Total chunks in Chroma DB: {total_chunks}")
                    
                    # Sample a few chunks but handle binary data safely
                    sample_docs = vector_store.similarity_search("test", k=3)
                    for i, doc in enumerate(sample_docs, 1):
                        # Use this more robust approach for potentially binary content
                        try:
                            # Try to print the content, but handle binary/non-UTF8 data
                            preview = repr(doc.page_content[:100])
                            logger.info(f"Sample Chunk {i} (preview): {preview}")
                        except Exception as e:
                            logger.warning(f"Could not preview chunk {i}: {e}")
                
                except Exception as e:
                    logger.error(f"Failed to verify Chroma DB: {e}")
                    
            verify_chroma_db(persist_directory)


            
            if progress_callback:
                progress_callback(f"✅ Successfully stored {len(all_documents)} chunks in Chroma DB")
        
        except Exception as e:
            logger.error(f"Failed to create Chroma DB: {e}")
            if progress_callback:
                progress_callback(f"❌ Failed to create Chroma DB: {e}")
    
    return {
        "successful_files": successful_files,
        "failed_files": failed_files,
        "status": "completed",
        "total_chunks": len(all_documents)
    }






def verify_embedding_model(embedding_model):
    """
    Test the embedding model to ensure it's working correctly
    
    Args:
        embedding_model: The embedding model to test
        
    Returns:
        bool: True if test is successful, False otherwise
    """
    try:
        # Use a simple test query
        test_text = "This is a test query to verify the embedding model is working"
        
        # Different methods for different embedding types
        if hasattr(embedding_model, 'embed_query'):
            embedding_result = embedding_model.embed_query(test_text)
        elif hasattr(embedding_model, 'embed_documents'):
            embedding_result = embedding_model.embed_documents([test_text])[0]
        else:
            # Most basic fallback
            try:
                embedding_result = embedding_model.embed(test_text)
            except Exception:
                return False
        
        # Check if the result looks reasonable
        if embedding_result and isinstance(embedding_result, list) and len(embedding_result) > 0:
            return True
        return False
    except Exception:
        return False

